"""
Export Engine — Generate translated documents in PDF, DOCX, and TXT formats.
Preserves structure: headings, paragraphs, bullet/numbered lists, and tables natively.
"""

from __future__ import annotations
import uuid
from pathlib import Path
from typing import List, Dict, Any, TYPE_CHECKING
from config import settings
from utils.logger import app_logger

if TYPE_CHECKING:
    from pipeline.reassembly import ReconstructedDocument, ReconstructedParagraph

def _ensure_export_dir() -> Path:
    settings.EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return settings.EXPORT_DIR


# ── Table Rendering Helpers ──────────────────────────────────────────────────

def _render_docx_table(doc: Any, cells: List[ReconstructedParagraph]) -> None:
    """Group table cells by row/col coordinates and build a native Word table."""
    rows_map: Dict[int, Dict[int, str]] = {}
    for c in cells:
        row_idx = c.metadata.get("row_index", 0)
        col_idx = c.metadata.get("col_index", 0)
        if row_idx not in rows_map:
            rows_map[row_idx] = {}
        rows_map[row_idx][col_idx] = c.translated_text
        
    if not rows_map:
        return
        
    num_rows = len(rows_map)
    num_cols = max(len(r) for r in rows_map.values()) if rows_map else 1
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"
    
    for r_idx in sorted(rows_map.keys()):
        for c_idx in sorted(rows_map[r_idx].keys()):
            if r_idx < num_rows and c_idx < num_cols:
                table.cell(r_idx, c_idx).text = rows_map[r_idx][c_idx]


def _render_pdf_table(story: list, cells: List[ReconstructedParagraph], body_style: Any) -> None:
    """Group table cells and render a native ReportLab table with wrapping paragraphs."""
    from reportlab.platypus import Table, TableStyle, Spacer
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    
    rows_map: Dict[int, Dict[int, str]] = {}
    for c in cells:
        row_idx = c.metadata.get("row_index", 0)
        col_idx = c.metadata.get("col_index", 0)
        if row_idx not in rows_map:
            rows_map[row_idx] = {}
        rows_map[row_idx][col_idx] = c.translated_text
        
    if not rows_map:
        return
        
    num_rows = len(rows_map)
    num_cols = max(len(r) for r in rows_map.values()) if rows_map else 1
    
    data = []
    from reportlab.platypus import Paragraph
    for r_idx in sorted(rows_map.keys()):
        row_data = []
        for c_idx in range(num_cols):
            text = rows_map[r_idx].get(c_idx, "")
            row_data.append(Paragraph(text, body_style))
        data.append(row_data)
        
    # Standard printable width on A4 page with 2cm margins is ~17cm
    col_width = 17.0 / num_cols
    table = Table(data, colWidths=[col_width * cm] * num_cols)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.4*cm))


# ── TXT Export ────────────────────────────────────────────────────────────────

def export_txt(
    job_id: str,
    source_text: str,
    translated_text: str,
    detected_language: str,
    confidence_score: float,
) -> Path:
    export_dir = _ensure_export_dir()
    out_path = export_dir / f"{job_id}_translated.txt"

    lang_label = {"ne": "Nepali", "si": "Sinhala", "en": "English"}.get(detected_language, "Unknown")

    content = (
        f"NeuroTranslate — Translation Export\n"
        f"{'='*50}\n"
        f"Source Language: {lang_label}\n"
        f"Confidence Score: {confidence_score:.1f}%\n"
        f"{'='*50}\n\n"
        f"ORIGINAL TEXT:\n{'-'*30}\n{source_text}\n\n"
        f"TRANSLATED TEXT (English):\n{'-'*30}\n{translated_text}\n"
    )

    out_path.write_text(content, encoding="utf-8")
    app_logger.info(f"TXT export: {out_path}")
    return out_path


# ── DOCX Export ───────────────────────────────────────────────────────────────

def export_docx(
    job_id: str,
    source_text: str,
    translated_text: str,
    detected_language: str,
    confidence_score: float,
    entities: List[dict] | None = None,
    reconstructed_doc: ReconstructedDocument | None = None,
) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    export_dir = _ensure_export_dir()
    out_path = export_dir / f"{job_id}_translated.docx"
    lang_label = {"ne": "Nepali", "si": "Sinhala", "en": "English"}.get(detected_language, "Unknown")

    doc = Document()

    # Title
    title = doc.add_heading("NeuroTranslate — Translation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = "Table Grid"
    meta_data = [
        ("Source Language", lang_label),
        ("Confidence Score", f"{confidence_score:.1f}%"),
        ("Model", "IndicTrans2" if detected_language == "ne" else "NLLB-200" if detected_language == "si" else "Pass-Through"),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v

    doc.add_paragraph()

    # Original text
    doc.add_heading("Original Text", level=1)
    for para in source_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_paragraph()

    # Translated text
    doc.add_heading("English Translation", level=1)

    if reconstructed_doc:
        accumulated_table_cells = []
        
        for page in reconstructed_doc.pages:
            for para in page.paragraphs:
                # Group cells by table index
                if para.block_type == "table_cell":
                    accumulated_table_cells.append(para)
                    continue
                else:
                    if accumulated_table_cells:
                        _render_docx_table(doc, accumulated_table_cells)
                        accumulated_table_cells = []
                
                # Render structural element
                if para.block_type == "heading":
                    doc.add_heading(para.translated_text, level=2)
                elif para.block_type == "list_item":
                    doc.add_paragraph(para.translated_text, style="List Bullet")
                elif para.block_type == "caption":
                    p = doc.add_paragraph(para.translated_text)
                    p.paragraph_format.left_indent = Pt(12)
                    p.runs[0].font.italic = True
                else:
                    doc.add_paragraph(para.translated_text)
            
            # Reconstruct any trailing table at page end
            if accumulated_table_cells:
                _render_docx_table(doc, accumulated_table_cells)
                accumulated_table_cells = []
    else:
        # Fallback to plain paragraphs splitting
        for para in translated_text.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # NER section
    if entities:
        doc.add_paragraph()
        doc.add_heading("Named Entities Detected", level=1)
        ner_table = doc.add_table(rows=1, cols=3)
        ner_table.style = "Table Grid"
        hdr = ner_table.rows[0].cells
        hdr[0].text = "Entity"
        hdr[1].text = "Type"
        hdr[2].text = "Confidence"
        for ent in entities[:50]:
            row = ner_table.add_row().cells
            row[0].text = ent.get("text", "")
            row[1].text = ent.get("label", "")
            row[2].text = f"{ent.get('score', 0) * 100:.1f}%"

    doc.save(str(out_path))
    app_logger.info(f"DOCX export: {out_path}")
    return out_path


# ── PDF Export ────────────────────────────────────────────────────────────────

def export_pdf(
    job_id: str,
    source_text: str,
    translated_text: str,
    detected_language: str,
    confidence_score: float,
    entities: List[dict] | None = None,
    reconstructed_doc: ReconstructedDocument | None = None,
) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak
    )

    export_dir = _ensure_export_dir()
    out_path = export_dir / f"{job_id}_translated.pdf"
    lang_label = {"ne": "Nepali", "si": "Sinhala", "en": "English"}.get(detected_language, "Unknown")

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title", parent=styles["Title"],
        fontSize=18, spaceAfter=16, textColor=colors.HexColor("#1a1a2e"),
    )
    heading_style = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontSize=13, spaceAfter=8, textColor=colors.HexColor("#16213e"),
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, spaceAfter=6, leading=14,
    )
    meta_style = ParagraphStyle(
        "Meta", parent=styles["Normal"],
        fontSize=9, textColor=colors.grey,
    )

    story = []

    # Title
    story.append(Paragraph("NeuroTranslate — Translation Report", title_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#6c63ff")))
    story.append(Spacer(1, 0.4*cm))

    # Metadata
    conf_color = "#22c55e" if confidence_score >= 90 else "#f59e0b" if confidence_score >= 70 else "#ef4444"
    model_used = "IndicTrans2" if detected_language == "ne" else "NLLB-200" if detected_language == "si" else "Pass-Through"
    meta_data = [
        ["Source Language", lang_label],
        ["Translation Model", model_used],
        ["Confidence Score", f"{confidence_score:.1f}%"],
    ]
    meta_table = Table(meta_data, colWidths=[5*cm, 12*cm])
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f4ff")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d7ff")),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor("#f9faff")]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 0.5*cm))

    # Original text
    story.append(Paragraph("Original Text", heading_style))
    for para in source_text.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip().replace("\n", "<br/>"), body_style))
    story.append(Spacer(1, 0.5*cm))

    # Translated text
    story.append(Paragraph("English Translation", heading_style))

    if reconstructed_doc:
        accumulated_table_cells = []
        
        for page_idx, page in enumerate(reconstructed_doc.pages):
            if page_idx > 0:
                story.append(PageBreak())
                
            for para in page.paragraphs:
                # Group cells by table index
                if para.block_type == "table_cell":
                    accumulated_table_cells.append(para)
                    continue
                else:
                    if accumulated_table_cells:
                        _render_pdf_table(story, accumulated_table_cells, body_style)
                        accumulated_table_cells = []
                
                # Render structural elements in PDF
                if para.block_type == "heading":
                    story.append(Paragraph(para.translated_text, heading_style))
                elif para.block_type == "list_item":
                    # Simple list styling
                    bullet_text = f"&bull; {para.translated_text}"
                    story.append(Paragraph(bullet_text, ParagraphStyle("List", parent=body_style, leftIndent=15)))
                elif para.block_type == "caption":
                    story.append(Paragraph(para.translated_text, ParagraphStyle("Cap", parent=body_style, fontName="Helvetica-Oblique", leftIndent=10)))
                else:
                    story.append(Paragraph(para.translated_text.replace("\n", "<br/>"), body_style))
            
            # Reconstruct any trailing table at page end
            if accumulated_table_cells:
                _render_pdf_table(story, accumulated_table_cells, body_style)
                accumulated_table_cells = []
    else:
        # Fallback to plain paragraphs splitting
        for para in translated_text.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip().replace("\n", "<br/>"), body_style))

    # NER section
    if entities:
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph("Named Entities Detected", heading_style))
        ner_data = [["Entity", "Type", "Score"]]
        for ent in entities[:30]:
            ner_data.append([ent.get("text",""), ent.get("label",""), f"{ent.get('score',0)*100:.1f}%"])
        ner_table = Table(ner_data, colWidths=[8*cm, 4*cm, 3*cm])
        ner_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#6c63ff")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.lightgrey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9faff")]),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(ner_table)

    doc.build(story)
    app_logger.info(f"PDF export: {out_path}")
    return out_path


# ── Dispatcher ────────────────────────────────────────────────────────────────

def export_all(job_id: str, result: dict, reconstructed_doc: ReconstructedDocument | None = None) -> dict:
    """Generate PDF, DOCX, and TXT exports. Returns dict of paths."""
    kwargs = dict(
        job_id=job_id,
        source_text=result.get("source_text", ""),
        translated_text=result.get("translated_text", ""),
        detected_language=result.get("detected_language", "en"),
        confidence_score=result.get("confidence_score", 0.0),
        entities=result.get("ner_entities"),
        reconstructed_doc=reconstructed_doc,
    )

    paths = {}
    try:
        paths["pdf"] = str(export_pdf(**kwargs))
    except Exception as e:
        app_logger.error(f"PDF export failed: {e}")
        paths["pdf"] = None

    try:
        paths["docx"] = str(export_docx(**kwargs))
    except Exception as e:
        app_logger.error(f"DOCX export failed: {e}")
        paths["docx"] = None

    try:
        paths["txt"] = str(export_txt(
            job_id=job_id,
            source_text=result.get("source_text", ""),
            translated_text=result.get("translated_text", ""),
            detected_language=result.get("detected_language", "en"),
            confidence_score=result.get("confidence_score", 0.0),
        ))
    except Exception as e:
        app_logger.error(f"TXT export failed: {e}")
        paths["txt"] = None

    return paths
