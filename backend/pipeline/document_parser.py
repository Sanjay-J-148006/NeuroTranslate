"""
Document parser — handles PDF, DOCX, and plain text extraction.
Priority 1: PDF (PyMuPDF)
Priority 2: DOCX (python-docx)
Priority 3: TXT/CSV (direct read)
"""

from __future__ import annotations
from pathlib import Path
from dataclasses import dataclass, field
from typing import List
from utils.logger import app_logger


@dataclass
class TextBlock:
    """A unit of extracted text with optional positional metadata."""
    text: str
    block_index: int
    page_number: int = 0
    block_type: str = "paragraph"   # paragraph / heading / table_cell / caption
    bbox: tuple | None = None        # (x0, y0, x1, y1) from OCR/PDF


@dataclass
class ParsedDocument:
    """Full parsed representation of an input document."""
    source_path: Path
    file_type: str
    blocks: List[TextBlock] = field(default_factory=list)
    needs_ocr: bool = False          # True if PDF is image-based / scanned


# ── PDF Parser ────────────────────────────────────────────────────────────────

def parse_pdf(file_path: Path) -> ParsedDocument:
    """Extract text from PDF using PyMuPDF. Detects scanned pages."""
    import fitz  # PyMuPDF

    doc = ParsedDocument(source_path=file_path, file_type="pdf")
    pdf = fitz.open(str(file_path))
    total_text_chars = 0

    for page_num, page in enumerate(pdf):
        page_text = page.get_text("blocks")  # list of (x0,y0,x1,y1,text,block_no,block_type)
        for i, block in enumerate(page_text):
            text = block[4].strip()
            if text:
                block_type = "heading" if block[6] == 0 and len(text) < 100 else "paragraph"
                doc.blocks.append(TextBlock(
                    text=text,
                    block_index=len(doc.blocks),
                    page_number=page_num,
                    block_type=block_type,
                    bbox=(block[0], block[1], block[2], block[3]),
                ))
                total_text_chars += len(text)

    pdf.close()

    # If very little text was extracted → likely scanned PDF → needs OCR
    if total_text_chars < 50:
        app_logger.info(f"PDF appears scanned (only {total_text_chars} chars) — flagging for OCR")
        doc.needs_ocr = True
        doc.blocks = []   # Clear; OCR engine will re-populate

    app_logger.info(f"PDF parsed: {len(doc.blocks)} blocks, {total_text_chars} chars, needs_ocr={doc.needs_ocr}")
    return doc


# ── DOCX Parser ───────────────────────────────────────────────────────────────

def parse_docx(file_path: Path) -> ParsedDocument:
    """Extract text blocks from DOCX preserving headings, paragraphs, tables."""
    from docx import Document

    doc = ParsedDocument(source_path=file_path, file_type="docx")
    docx = Document(str(file_path))

    for para in docx.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name.lower() if para.style else ""
        block_type = "heading" if "heading" in style_name else "paragraph"
        doc.blocks.append(TextBlock(
            text=text,
            block_index=len(doc.blocks),
            block_type=block_type,
        ))

    # Extract table cells
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    doc.blocks.append(TextBlock(
                        text=text,
                        block_index=len(doc.blocks),
                        block_type="table_cell",
                    ))

    app_logger.info(f"DOCX parsed: {len(doc.blocks)} blocks")
    return doc


# ── TXT / CSV Parser ─────────────────────────────────────────────────────────

def parse_text(file_path: Path) -> ParsedDocument:
    """Read plain text or CSV as paragraph blocks."""
    doc = ParsedDocument(source_path=file_path, file_type="txt")

    content = file_path.read_text(encoding="utf-8", errors="replace")
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [line.strip() for line in content.splitlines() if line.strip()]

    for i, para in enumerate(paragraphs):
        doc.blocks.append(TextBlock(text=para, block_index=i))

    app_logger.info(f"Text parsed: {len(doc.blocks)} blocks")
    return doc


# ── Dispatcher ────────────────────────────────────────────────────────────────

def parse_document(file_path: Path, file_type: str) -> ParsedDocument:
    """Route file to the correct parser based on type category."""
    file_path = Path(file_path)

    if file_type == "pdf":
        return parse_pdf(file_path)
    elif file_type == "docx":
        return parse_docx(file_path)
    elif file_type == "txt":
        return parse_text(file_path)
    elif file_type in ("image", "audio", "video"):
        # Images/audio/video are handled by OCR or ASR engine directly
        return ParsedDocument(source_path=file_path, file_type=file_type, needs_ocr=True)
    else:
        raise ValueError(f"Unsupported file type for document parser: {file_type}")
